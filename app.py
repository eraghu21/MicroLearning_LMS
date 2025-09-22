import streamlit as st
import pandas as pd
import pyAesCrypt
import os
import json
import datetime
from io import BytesIO
from zipfile import ZipFile
from email.message import EmailMessage
import smtplib
from docx import Document

# -------------------- Configuration --------------------
BUFFER_SIZE = 64 * 1024
AES_FILE = "Students_List.xlsx.aes"   # updated filename
PROGRESS_FILE = "progress.json"
YOUTUBE_VIDEO_URL = "https://youtu.be/Tva_sr4BUfk?si=GukUUa-tY-VvYo73"  # Replace with your video URL

# -------------------- Load Secrets --------------------
# Fallback values if secrets are missing
password = None
email_sender = None
email_password = None

try:
    password = st.secrets["encryption"]["password"]
    email_sender = st.secrets["email"]["sender"]
    email_password = st.secrets["email"]["password"]
except Exception:
    # For local testing (⚠️ replace with your real values)
    password = "yourpassword"
    email_sender = "youremail@gmail.com"
    email_password = "your-email-app-password"

# -------------------- Helper Functions --------------------
@st.cache_data(show_spinner=False)
def load_students_list():
    try:
        # Try AES decryption
        with open(AES_FILE, "rb") as fIn:
            decrypted = BytesIO()
            pyAesCrypt.decryptStream(
                fIn, decrypted, password, BUFFER_SIZE, os.path.getsize(AES_FILE)
            )
            decrypted.seek(0)
            df = pd.read_excel(decrypted)
    except Exception:
        # Fallback: try plain Excel if not AES
        df = pd.read_excel(AES_FILE)

    # 🔑 Ensure RegNo is always string for comparison
    df["RegNo"] = df["RegNo"].astype(str).str.strip()
    return df

def load_progress():
    if not os.path.exists(PROGRESS_FILE):
        return {}
    with open(PROGRESS_FILE, "r") as f:
        return json.load(f)

def save_progress(progress):
    with open(PROGRESS_FILE, "w") as f:
        json.dump(progress, f, indent=4)

def generate_certificate(name, regno):
    doc = Document()
    doc.add_heading("Certificate of Completion", 0)
    doc.add_paragraph(f"This is to certify that {name} ({regno}) has successfully completed the video.")
    doc.add_paragraph(f"Date: {datetime.date.today()}")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def send_email(to, name, regno, cert_bytes):
    if not email_sender or not email_password:
        st.warning("⚠️ Email credentials not configured. Skipping email send.")
        return
    msg = EmailMessage()
    msg['Subject'] = "Your Certificate of Completion"
    msg['From'] = email_sender
    msg['To'] = to
    msg.set_content(f"Dear {name},\n\nCongratulations! Your certificate is attached.\n\nRegards,\nAdmin")
    msg.add_attachment(cert_bytes, maintype='application',
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename=f"{regno}_certificate.docx")
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(email_sender, email_password)
        smtp.send_message(msg)

# -------------------- Streamlit App --------------------
st.title("🎓 Microlearning Certificate Portal")

regno = st.text_input("Enter your Registration Number").strip().upper()

if regno:
    students_df = load_students_list()
    student = students_df[students_df["RegNo"] == regno]
    if not student.empty:
        name = student.iloc[0]["Name"]
        email = student.iloc[0]["Email"]
        st.success(f"Welcome, {name}!")
        progress = load_progress()
        record = progress.get(regno, {})
        
        if record.get("video_completed"):
            st.info("✅ You have already completed the video.")
            if st.button("Download Certificate Again"):
                cert_bytes = generate_certificate(name, regno)
                st.download_button("⬇️ Download Certificate", cert_bytes, file_name=f"{regno}_certificate.docx")
        else:
            st.video(YOUTUBE_VIDEO_URL)
            if st.button("I have watched the complete video"):
                cert_bytes = generate_certificate(name, regno)
                st.success("🎉 Video marked as complete. Your certificate is ready!")
                st.download_button("⬇️ Download Certificate", cert_bytes, file_name=f"{regno}_certificate.docx")
                try:
                    send_email(email, name, regno, cert_bytes)
                    st.success(f"📩 Certificate sent to {email}")
                except:
                    st.warning("⚠️ Failed to send email.")
                progress[regno] = {
                    "name": name,
                    "email": email,
                    "video_completed": True,
                    "certificate_sent": True,
                    "timestamp": str(datetime.datetime.now())
                }
                save_progress(progress)
    else:
        st.error("Registration number not found.")
